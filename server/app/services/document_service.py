"""文档解析服务（Task A / V2.0 §9.7）

支持 PDF / Word / Markdown / TXT 格式解析和切片。
"""

import logging
import uuid
from pathlib import Path

from sqlalchemy.orm import Session

from app.llm.provider import get_llm
from app.models.document import Document, Segment
from app.vectors.base import create_vector_store

logger = logging.getLogger("bluelink.document")

# 切片参数
MAX_CHUNK_SIZE = 1024       # 目标切片长度（字符）
CHUNK_OVERLAP = 256          # 相邻切片重叠部分（前后各 128，合计 256）
MERGE_THRESHOLD = 80         # 小于此长度的段落与下一段合并


class DocumentService:

    @staticmethod
    def parse_and_chunk(file_path: str) -> list[dict]:
        """解析文档文件，按段落切片

        Args:
            file_path: 文件路径

        Returns:
            [{"index": int, "text": str}, ...]
        """
        ext = Path(file_path).suffix.lower()

        if ext == ".pdf":
            text = DocumentService._extract_pdf(file_path)
        elif ext in (".docx", ".doc"):
            text = DocumentService._extract_docx(file_path)
        else:
            # .md, .txt 等纯文本
            with open(file_path, encoding="utf-8", errors="replace") as f:
                text = f.read()

        return DocumentService._chunk_by_paragraph(text)

    @staticmethod
    def _extract_pdf(file_path: str) -> str:
        """PyMuPDF 提取 PDF 文本"""
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        pages = []
        for page in doc:
            text = page.get_text().strip()
            if text:
                pages.append(text)
        doc.close()
        return "\n\n".join(pages)

    @staticmethod
    def _extract_docx(file_path: str) -> str:
        """python-docx 提取 Word 文本"""
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)

    @staticmethod
    def _chunk_by_paragraph(
        text: str,
        max_len: int = MAX_CHUNK_SIZE,
        overlap: int = CHUNK_OVERLAP,
        merge_threshold: int = MERGE_THRESHOLD,
    ) -> list[dict]:
        """滑动窗口切片，支持重叠

        1. 按空行分段
        2. 过短的段落与下一段合并
        3. 目标 ≈ max_len 字符
        4. 相邻切片之间保留 overlap 字符作为上下文

        Args:
            text: 全文
            max_len: 目标切片长度
            overlap: 前后重叠字符数
            merge_threshold: 短段落合并阈值

        Returns:
            [{"index": int, "text": str}, ...]
        """
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]

        # 合并短段落
        merged: list[str] = []
        for p in paragraphs:
            if merged and len(p) < merge_threshold:
                merged[-1] += "\n" + p
            else:
                merged.append(p)

        # 滑动窗口切片
        chunks: list[dict] = []
        idx = 0
        buf = ""

        for p in merged:
            # 加上当前段落后超过上限且已有内容 → 输出当前切片
            if len(buf) + len(p) > max_len and buf:
                chunks.append({"index": idx, "text": buf.strip()})
                idx += 1
                # 保留尾部 overlap 字符作为下一个切片的开头上下文
                buf = DocumentService._tail_context(buf, overlap) + "\n" + p + "\n"
            else:
                buf += p + "\n"

        if buf.strip():
            chunks.append({"index": idx, "text": buf.strip()})

        return chunks

    @staticmethod
    def _tail_context(text: str, overlap: int) -> str:
        """取文本尾部 overlap 字符（在段落边界处截断）"""
        if len(text) <= overlap or overlap <= 0:
            return ""
        tail = text[-overlap:]
        nl = tail.find("\n")
        if nl != -1:
            return tail[nl + 1:]
        return tail

    @staticmethod
    def save_document(
        db: Session,
        user_id: str,
        title: str,
        source: str | None,
        chunks: list[dict],
        privacy_level: str = "CLOUD_OK",
    ) -> Document:
        """保存文档及切片到数据库，并向量化

        Args:
            db: 数据库会话
            user_id: 用户 ID
            title: 文档标题
            source: 来源路径
            chunks: parse_and_chunk 的输出
            privacy_level: 隐私等级

        Returns:
            创建的 Document ORM 对象
        """
        import time
        now = int(time.time() * 1000)
        doc_id = str(uuid.uuid4())

        # 1. 创建文档记录
        doc = Document(
            id=doc_id,
            user_id=user_id,
            title=title,
            privacy_level=privacy_level,
            source=source,
            created_at=now,
        )
        db.add(doc)
        db.flush()

        # 2. 创建切片记录
        segments: list[Segment] = []
        for chunk in chunks:
            seg = Segment(
                id=str(uuid.uuid4()),
                doc_id=doc_id,
                index_num=chunk["index"],
                text=chunk["text"],
                summary=chunk["text"][:80],
            )
            segments.append(seg)
            db.add(seg)
        db.flush()

        # 3. 先提交 DB 事务，确保 segments 已持久化（否则 VectorStore 外键约束会失败）
        db.commit()

        # 4. 向量化并存入 VectorStore（分批，每次 ≤ 10 条）
        try:
            llm = get_llm()
            store = create_vector_store()
            batch_size = 10
            total_vectors = 0
            for i in range(0, len(segments), batch_size):
                batch = segments[i:i + batch_size]
                texts = [s.text for s in batch]
                vectors = llm.embed_texts(texts)
                if len(vectors) != len(batch):
                    logger.warning(
                        "向量数量不匹配: 期望 %d, 收到 %d（跳过该批）",
                        len(batch), len(vectors),
                    )
                    continue
                segment_vectors = [
                    (s.id, vec) for s, vec in zip(batch, vectors)
                ]
                store.save_segment_vectors(doc_id, segment_vectors)
                total_vectors += len(vectors)
            logger.info("向量化成功: %d 个切片", total_vectors)
        except Exception as e:
            logger.warning("向量化失败（跳过）: %s", e)
        logger.info(
            "文档保存成功 user=%s doc=%s chunks=%d",
            user_id, doc_id, len(segments),
        )
        return doc
